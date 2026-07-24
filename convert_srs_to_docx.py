"""
Convert SRS markdown to Word document (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_title(doc, text):
    """Add a title to the document"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 42, 74)  # Navy blue
    return title

def add_subtitle(doc, text):
    """Add a subtitle"""
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(46, 134, 171)  # Teal
    return subtitle

def add_metadata(doc, items):
    """Add metadata section"""
    for label, value in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.font.bold = True
        p.add_run(value)

def parse_and_add_content(doc, markdown_text):
    """Parse markdown and add to document"""
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Main title (# Title)
        if line.startswith('# ') and not line.startswith('# '):
            text = line[2:].strip()
            add_title(doc, text)
            i += 1
            continue
        
        # Subtitle (## Subtitle)
        if line.startswith('## '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=1)
            heading_format = heading.paragraph_format
            heading_format.space_before = Pt(12)
            heading_format.space_after = Pt(6)
            i += 1
            continue
        
        # Subheading (### Subheading)
        if line.startswith('### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=2)
            heading_format = heading.paragraph_format
            heading_format.space_before = Pt(6)
            heading_format.space_after = Pt(3)
            i += 1
            continue
        
        # Subsubheading (#### Subsubheading)
        if line.startswith('#### '):
            text = line[5:].strip()
            heading = doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Horizontal rule (---)
        if line.strip() == '---':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Bullet list
        if line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:].strip()
            
            # Parse inline formatting
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line).strip()
            add_formatted_text(p, text)
            i += 1
            continue
        
        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())
            i += 1
            continue
        
        i += 1

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code)"""
    # Split by formatting markers
    parts = re.split(r'(\*\*.*?\*\*|_.*?_|`.*?`)', text)
    
    for part in parts:
        if not part:
            continue
        
        # Bold text
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        # Italic text
        elif part.startswith('_') and part.endswith('_'):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        # Code/monospace
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            run._element.get_or_add_rPr().append(shading_elm)
        # Regular text
        else:
            paragraph.add_run(part)

def main():
    # Read markdown SRS
    with open('docs/VOO_Ward_SRS.md', 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse markdown and add to document
    parse_and_add_content(doc, markdown_content)
    
    # Save document
    output_path = 'docs/VOO_Ward_SRS.docx'
    doc.save(output_path)
    print(f"✓ SRS converted to: {output_path}")

if __name__ == '__main__':
    main()
