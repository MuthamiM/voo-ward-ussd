#!/usr/bin/env python3
"""
Convert VOO Ward SRS markdown to Word document (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

def convert_md_to_docx(md_file, output_file):
    """Convert markdown file to Word document"""
    
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Software Requirements Specification (SRS)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('VOO Ward Citizen Engagement Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.bold = True
    
    # Split by lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            doc.add_heading(heading, level=3)
        elif line.startswith('# '):
            heading = line[2:].strip()
            doc.add_heading(heading, level=1)
        
        # Handle horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()
        
        # Handle bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Process inline formatting
            text = format_text(text)
            p = doc.add_paragraph(text, style='List Bullet')
            apply_formatting(p, text)
        
        # Regular paragraphs
        elif line.strip():
            text = line.strip()
            # Skip markdown formatting markers in regular text processing
            if not text.startswith('#') and not text.startswith('---'):
                text = format_text(text)
                p = doc.add_paragraph(text)
                apply_formatting(p, text)
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✓ Converted to: {output_file}")

def format_text(text):
    """Format inline markdown syntax"""
    # This is a simple version; for production consider a markdown parser
    return text

def apply_formatting(paragraph, text):
    """Apply bold, italic formatting based on markdown markers"""
    # Clear existing runs
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    
    # Re-add with formatting
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)', text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('__') and part.endswith('__'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)

if __name__ == "__main__":
    md_file = r"c:\Users\kivuv\Documents\voo-ward-ussd\docs\VOO_Ward_SRS.md"
    output_file = r"c:\Users\kivuv\Documents\voo-ward-ussd\docs\VOO_Ward_SRS.docx"
    
    if os.path.exists(md_file):
        convert_md_to_docx(md_file, output_file)
    else:
        print(f"Error: {md_file} not found")
