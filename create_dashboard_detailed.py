#!/usr/bin/env python3
"""
Create a detailed, professional Word document from AdminDashboard presentation
With proper sections, formatting, and structure
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_heading_style(para, level, text):
    """Set proper heading style"""
    para.text = text
    para.style = f'Heading {level}'
    return para

def add_section_break(doc):
    """Add a section break"""
    doc.add_page_break()

def format_bullet_list(doc, items):
    """Add a formatted bullet list"""
    for item in items:
        if item.strip():
            p = doc.add_paragraph(item.strip(), style='List Bullet')
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

def format_numbered_list(doc, items):
    """Add a formatted numbered list"""
    for item in items:
        if item.strip():
            p = doc.add_paragraph(item.strip(), style='List Number')

def create_detailed_docx():
    """Create a detailed professional Word document for Admin Dashboard presentation"""
    
    doc = Document()
    
    # Set default styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ===== SLIDE 1: TITLE PAGE =====
    title = doc.add_heading('VOO Ward Admin Dashboard', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Ward Management & Citizen Engagement Platform')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.bold = True
    subtitle_format.size = Pt(14)
    
    tagline = doc.add_paragraph('Kyamatu Ward — Modern Governance')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_format = tagline.runs[0]
    tagline_format.italic = True
    tagline_format.size = Pt(12)
    
    doc.add_paragraph()  # Spacing
    
    # Key features
    features = [
        'Issue Tracking & Resolution',
        'Bursary Application Management',
        'Real-Time Analytics & Maps',
        'AI Chat Assistant (Mai)'
    ]
    for feature in features:
        p = doc.add_paragraph(feature)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Presented by
    presented = doc.add_paragraph('Presented by: VOO Ward Development Team')
    presented.alignment = WD_ALIGN_PARAGRAPH.CENTER
    presented_format = presented.runs[0]
    presented_format.size = Pt(10)
    
    add_section_break(doc)
    
    # ===== SLIDE 2: TABLE OF CONTENTS =====
    toc_heading = doc.add_heading('Table of Contents', 1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_items = [
        '01. Project Overview',
        '02. Dashboard Overview',
        '03. Issue Management',
        '04. Bursary Management',
        '05. Announcements & USSD',
        '06. Constituent Registry',
        '07. User Management & Roles',
        '08. Analytics & Visualization',
        '09. Security & Authentication',
        '10. AI Assistant & Notifications',
        '11. Technology Stack',
        '12. Summary & Roadmap'
    ]
    format_numbered_list(doc, toc_items)
    
    add_section_break(doc)
    
    # ===== SLIDE 3: PROJECT OVERVIEW =====
    doc.add_heading('01. Project Overview', 1)
    doc.add_heading('Kyamatu Ward — Comprehensive Citizen Engagement Platform', 2)
    
    overview = doc.add_paragraph(
        'VOO Ward is a full-stack governance platform enabling the Member of County Assembly (MCA) '
        'and staff to manage ward operations, track citizen issues, process bursary applications, '
        'and communicate with constituents through multiple channels including web dashboard, USSD, and mobile app.'
    )
    
    doc.add_heading('Dashboard Modules', 3)
    modules = [
        'Web Admin Dashboard — Full management interface',
        'USSD Service (*384#) — Feature phone access',
        'Mobile App (React Native) — Citizen portal',
        'WhatsApp Integration — Messaging channel',
        'SMS Notifications — Announcement delivery',
        'AI Chat Assistant — In-dashboard help'
    ]
    format_bullet_list(doc, modules)
    
    doc.add_heading('Core Objectives', 3)
    objectives = [
        'Transparent governance & issue resolution',
        'Efficient bursary application processing',
        'Real-time communication with constituents'
    ]
    format_bullet_list(doc, objectives)
    
    add_section_break(doc)
    
    # ===== SLIDE 4: DASHBOARD OVERVIEW =====
    doc.add_heading('02. Dashboard Overview', 1)
    doc.add_heading('Layout, Navigation & Live Statistics', 2)
    
    doc.add_heading('Sidebar Navigation', 3)
    nav_items = [
        '🏠  Dashboard',
        '⚠️  Issues',
        '🎓  Bursaries',
        '👥  Constituents',
        '📢  Announcements',
        '📱  USSD Interactions',
        '⚙️  Settings'
    ]
    format_bullet_list(doc, nav_items)
    
    doc.add_heading('Dashboard Features', 3)
    features = [
        'Animated stat counters — values count up from 0 on page load',
        'Glassmorphism UI — backdrop blur with semi-transparent cards',
        'Real-time updates — 30-second polling for new data',
        'Breadcrumb navigation — contextual page location',
        'Responsive design — adapts to desktop, tablet & mobile viewports'
    ]
    format_bullet_list(doc, features)
    
    add_section_break(doc)
    
    # ===== SLIDE 5: ISSUE MANAGEMENT =====
    doc.add_heading('03. Issue Management', 1)
    doc.add_heading('Citizen Issue Tracking, Resolution & Reporting', 2)
    
    doc.add_heading('Core Capabilities', 3)
    capabilities = [
        'View all issues — ticket ID, category, message, phone, status, source',
        'Filter by status: All / Pending / In Progress / Resolved / Closed',
        'Add new issue via modal (title, description, priority, location)',
        'Update status with action notes (who acted, when)',
        'Bulk resolve multiple issues (MCA only, max 200/batch)',
        'Categories: Infrastructure, Water, Health, & custom',
        'Auto-update USSD interactions on resolution',
        'Status summary footer with pending/in-progress/resolved counts'
    ]
    format_bullet_list(doc, capabilities)
    
    doc.add_heading('Status Workflow', 3)
    workflow = doc.add_paragraph('Open → In Progress → Resolved → Closed')
    workflow_items = [
        'Each transition recorded with timestamp & user',
        'Action notes required for status changes'
    ]
    format_bullet_list(doc, workflow_items)
    
    doc.add_heading('Data Sources', 3)
    sources = [
        'Web dashboard — manual issue creation',
        'USSD — citizen-reported via feature phones',
        'Mobile App — photo & GPS-tagged reports'
    ]
    format_bullet_list(doc, sources)
    
    doc.add_heading('Priority Levels', 3)
    priorities = [
        '🔴  Urgent — immediate attention required',
        '🟠  High — address within 24 hours',
        '🟡  Medium — standard processing',
        '⚪  Low — when resources permit'
    ]
    format_bullet_list(doc, priorities)
    
    add_section_break(doc)
    
    # ===== SLIDE 6: BURSARY MANAGEMENT =====
    doc.add_heading('04. Bursary Management', 1)
    doc.add_heading('Student Bursary Application Processing & Disbursement', 2)
    
    doc.add_heading('Application Processing', 3)
    app_proc = [
        'View all bursary applications (ref code, name, institution, amount)',
        'Process via modal: applicant, ID, amount (KES), decision, notes',
        'MCA-only access for viewing and status updates',
        'Export bursary data to CSV',
        'Linked audit trail for all bursary actions'
    ]
    format_bullet_list(doc, app_proc)
    
    doc.add_heading('Status Workflow', 3)
    burse_workflow = [
        'Pending → Under Review → Approved → Disbursed',
        'Pending → Under Review → Rejected (with reason)',
        'Each status change recorded with timestamp',
        'Decision notes and reason tracking',
        'Amount tracking in Kenyan Shillings (KES)'
    ]
    format_bullet_list(doc, burse_workflow)
    
    doc.add_heading('Bursary Analytics', 3)
    analytics = [
        'Bar chart: applications by status (Pending/Approved/Rejected/Received)',
        'Total applications count',
        'Total amount requested vs. approved amount',
        'Approval rate percentage',
        'Filterable by date range',
        'Animated counters for key metrics'
    ]
    format_bullet_list(doc, analytics)
    
    doc.add_heading('Access Control', 3)
    access = [
        'Only MCA (Full Access) role can:',
        '    — View bursary applications',
        '    — Approve or reject applications',
        '    — Modify disbursement status',
        '    — Export bursary reports',
        'PA/Clerk roles have no bursary access'
    ]
    format_bullet_list(doc, access)
    
    add_section_break(doc)
    
    # ===== SLIDE 7: ANNOUNCEMENTS & USSD =====
    doc.add_heading('05. Announcements & USSD', 1)
    doc.add_heading('Ward Communication & Multi-Channel Outreach', 2)
    
    doc.add_heading('Announcements System', 3)
    announce = [
        'Create announcements with type selection:',
        '    — General Information',
        '    — Urgent Notice',
        '    — Event Notification',
        '    — Service Update',
        'Optional SMS notification on publish',
        'Delete old announcements',
        'Timestamped creation date display'
    ]
    format_bullet_list(doc, announce)
    
    doc.add_heading('USSD Integration (*384#)', 3)
    ussd_items = [
        'Full USSD menu system for feature phones',
        'USSD News Preview — monospace text preview',
        'News feed API combines announcements + resolved issues',
        'Paginated display optimized for USSD screens',
        'Session logging: phone, input, response, IP, timestamp',
        'USSD interaction export to CSV'
    ]
    format_bullet_list(doc, ussd_items)
    
    doc.add_heading('Citizen Messages', 3)
    messages = [
        'Public message submission (no auth required)',
        'Fields: name, phone, subject, message',
        'Status workflow: Unread → Read → Replied',
        'Admin can reply directly from dashboard'
    ]
    format_bullet_list(doc, messages)
    
    doc.add_heading('USSD Menu Structure', 3)
    ussd_menu = doc.add_paragraph()
    ussd_menu.add_run('1. Register as Voter\n2. Report an Issue\n3. Check Issue Status\n4. My Registration\n5. Announcements\n6. Contact Us\n0. Exit')
    
    add_section_break(doc)
    
    # ===== SLIDE 8: CONSTITUENT REGISTRY =====
    doc.add_heading('06. Constituent Registry', 1)
    doc.add_heading('Voter Registration & Citizen Data Management', 2)
    
    doc.add_heading('Registry Features', 3)
    registry = [
        'View all registered constituents in table view',
        'Data fields: phone, national ID, full name, location',
        'Sub-location and village-level grouping',
        'Verification status tracking per constituent',
        'Verify or reject with reason documentation',
        'Export full registry to CSV'
    ]
    format_bullet_list(doc, registry)
    
    doc.add_heading('Registration Channels', 3)
    channels = [
        'Web admin — manual registration',
        'USSD — self-registration via *384#',
        'Mobile App — 5-step registration with ID & selfie',
        'ID document photo capture',
        'Selfie verification',
        'Polling station assignment'
    ]
    format_bullet_list(doc, channels)
    
    doc.add_heading('Data Management', 3)
    data_mgmt = [
        'Table sorting — click column headers (numeric-aware)',
        'Sort direction indicators (▲/▼)',
        'Pagination: configurable 10/20/50/100 per page',
        'Search and filter capabilities',
        'Secure data handling with input sanitization'
    ]
    format_bullet_list(doc, data_mgmt)
    
    doc.add_heading('Verification Workflow', 3)
    verification = [
        'Citizens register via any channel',
        'Status tracked: Pending / Verified / Rejected',
        'Admin reviews and verifies identity',
        'Rejection requires documented reason',
        'Verified citizens can access all services'
    ]
    format_bullet_list(doc, verification)
    
    add_section_break(doc)
    
    # ===== SLIDE 9: USER MANAGEMENT & ROLES =====
    doc.add_heading('07. User Management & Roles', 1)
    doc.add_heading('Admin Accounts, Access Control & Permissions', 2)
    
    doc.add_heading('MCA — Full Access', 3)
    mca_perms = [
        'Super Admin / Member of County Assembly',
        'Full access to all modules',
        'Manage users (add/edit/delete)',
        'Process bursary applications',
        'Bulk issue resolution (max 200)',
        'Export all data types',
        'Edit AI chatbot knowledge base',
        'Protected account — cannot be deleted'
    ]
    format_bullet_list(doc, mca_perms)
    
    doc.add_heading('PA — Personal Assistant', 3)
    pa_perms = [
        'Assigned by MCA (max 2 PA/Clerk)',
        'Issue management & status updates',
        'View and create announcements',
        'Export issues & USSD data',
        'Activity & audit log access',
        'No bursary access',
        'No user management access'
    ]
    format_bullet_list(doc, pa_perms)
    
    doc.add_heading('Clerk — Basic Operational Access', 3)
    clerk_perms = [
        'Basic operational access',
        'Issue viewing and management',
        'Announcement creation',
        'Limited dashboard access',
        'Self-registration via signup page',
        'No bursary or user management',
        'No export capabilities'
    ]
    format_bullet_list(doc, clerk_perms)
    
    doc.add_heading('User Management Features', 3)
    user_features = doc.add_paragraph(
        'Card grid view with stats (resolved/active issues, logins) • '
        'Add via signup page • Edit & reset passwords • '
        'Max 3 total users (1 MCA + 2 PA/Clerk) • '
        'Filter by search & role • Last login tracking'
    )
    
    add_section_break(doc)
    
    # ===== SLIDE 10: ANALYTICS & VISUALIZATION =====
    doc.add_heading('08. Analytics & Visualization', 1)
    doc.add_heading('Charts, Interactive Maps & Data Insights', 2)
    
    doc.add_heading('📊  Issue Status (Doughnut)', 3)
    issue_status = [
        'Pending / In Progress / Resolved / Closed',
        '60% cutout doughnut chart',
        'Color-coded segments',
        'Interactive tooltips',
        'Powered by Chart.js'
    ]
    format_bullet_list(doc, issue_status)
    
    doc.add_heading('📈  Issues Over Time (Line)', 3)
    issues_time = [
        'Dual dataset: New vs Resolved',
        '30-day rolling window',
        'Trend analysis capability',
        'Responsive canvas rendering',
        'Date range filtering'
    ]
    format_bullet_list(doc, issues_time)
    
    doc.add_heading('📊  Bursary Status (Bar)', 3)
    burse_status = [
        'Pending / Approved / Rejected / Received',
        'Status distribution bar chart',
        'Amount tracking per status',
        'MCA-only analytics view',
        'Exportable data'
    ]
    format_bullet_list(doc, burse_status)
    
    doc.add_heading('🗺️  Interactive Issue Map (Leaflet.js)', 3)
    map_features = [
        'OpenStreetMap tiles centered on Kyamatu Ward',
        'Color-coded markers: Red=Pending, Amber=In Progress, Green=Resolved',
        'Custom markers with priority indicators (! / ● / ○)',
        'Interactive popups with issue details & \'View Details\' button',
        'Map statistics panel: Total, Pending, In Progress, Resolved',
        'Filter by status and priority'
    ]
    format_bullet_list(doc, map_features)
    
    doc.add_heading('📅  Date Range Filters', 3)
    date_filters = [
        'All Time — full data view',
        'Today — current day activity',
        'Last 7 Days — weekly report',
        'Last 30 Days — monthly summary',
        'Last 90 Days — quarterly review',
        'Custom Range — pick start & end dates'
    ]
    format_bullet_list(doc, date_filters)
    
    add_section_break(doc)
    
    # ===== SLIDE 11: SECURITY & AUTHENTICATION =====
    doc.add_heading('09. Security & Authentication', 1)
    doc.add_heading('Multi-Layer Protection & Access Control', 2)
    
    doc.add_heading('🔐  Authentication', 3)
    auth = [
        'Username + Password login',
        'Glassmorphism login card UI',
        'Bearer token session management',
        '30-minute session timeout',
        '24-hour client-side token expiry',
        'Single session enforcement',
        'Signup page for new PA/Clerk'
    ]
    format_bullet_list(doc, auth)
    
    doc.add_heading('🛡️  Protection Measures', 3)
    protection = [
        'bcrypt password hashing (auto-migrates SHA-256)',
        'Rate limiting: 10 attempts/min/IP',
        'Client-side: 5 attempts → 5-min lockout',
        'Input sanitization (XSS, injection)',
        'Input validation (server + client)',
        'CORS whitelisted origins only',
        'Trust proxy for production (Render)'
    ]
    format_bullet_list(doc, protection)
    
    doc.add_heading('📋  Security Headers', 3)
    headers = [
        'X-Frame-Options: DENY',
        'X-Content-Type-Options: nosniff',
        'X-XSS-Protection: 1; mode=block',
        'Referrer-Policy: strict-origin',
        'Content-Security-Policy (CSP)',
        'Cache-Control: no-store',
        'HTTPS enforced in production'
    ]
    format_bullet_list(doc, headers)
    
    doc.add_heading('📝  Audit Trail', 3)
    audit = [
        'Timeline view of all admin actions',
        'Action types: CREATE, UPDATE, DELETE, LOGIN, LOGOUT, EXPORT',
        'Logged: user, role, IP address, timestamp, description',
        'Filter by action type, user, date, search term',
        'Pagination (20/page) & CSV export'
    ]
    format_bullet_list(doc, audit)
    
    doc.add_heading('🔑  Password Management', 3)
    password = [
        'Settings page: change password (verify current first)',
        'Admin can reset PA/Clerk passwords',
        'Minimum 6 characters, max 128',
        'bcrypt with salt rounds = 10',
        'Legacy hash auto-migration on login'
    ]
    format_bullet_list(doc, password)
    
    add_section_break(doc)
    
    # ===== SLIDE 12: AI ASSISTANT & NOTIFICATIONS =====
    doc.add_heading('10. AI Assistant & Notifications', 1)
    doc.add_heading('Smart Features, Real-Time Updates & Alerts', 2)
    
    doc.add_heading('🤖  AI Chat Assistant (Mai)', 3)
    ai_features = [
        'Floating chat button (bottom-right, animated bounce)',
        'Chat modal with full conversation history',
        'Contextual responses: navigation, issues, exports, passwords',
        'Action buttons deep-link to relevant dashboard sections',
        'Database connectivity check from chat',
        'Knowledge base topics: business, tech, health, education, science',
        'MCA can edit chatbot knowledge base via API',
        'Appears automatically after login'
    ]
    format_bullet_list(doc, ai_features)
    
    doc.add_heading('🔔  Notification System', 3)
    notif = [
        'Header bell icon with unread badge (pulsing animation)',
        'Dropdown panel (380px) with scrollable list',
        'Types: Issue ⚠️, Bursary 💰, User 👤, System ℹ️, Success ✅',
        'Mark individual or all as read',
        '30-second polling for new notifications',
        'localStorage persistence (max 50 items)',
        'Time-ago formatting (Just now, Xm, Xh, Xd)',
        'Desktop browser notifications (Notification API)'
    ]
    format_bullet_list(doc, notif)
    
    doc.add_heading('📡  Real-Time Updates', 3)
    realtime = [
        '30-second polling for new issues and data changes',
        'Auto-refresh: issues table, announcements, map, analytics',
        'Notification sound (WAV audio) on new issue',
        'Event system (on/off/emit) for cross-module communication',
        'Desktop Notification API integration with permission request'
    ]
    format_bullet_list(doc, realtime)
    
    doc.add_heading('🟢  Connection Status', 3)
    connection = [
        'Fixed bottom-right indicator dot',
        'Green = Live, Amber = Connecting, Red = Offline',
        'Automatic reconnection on failure',
        'Graceful degradation when API unreachable'
    ]
    format_bullet_list(doc, connection)
    
    add_section_break(doc)
    
    # ===== SLIDE 13: TECHNOLOGY STACK =====
    doc.add_heading('11. Technology Stack', 1)
    doc.add_heading('Architecture, Libraries & Deployment Infrastructure', 2)
    
    doc.add_heading('⚙️  Backend', 3)
    backend = [
        'Node.js + Express.js',
        'PostgreSQL (Supabase) — primary DB',
        'MongoDB Atlas — secondary store',
        'Redis — caching layer',
        'bcryptjs — password security',
        'multer + sharp — image processing',
        'AWS S3 v3 — cloud storage (optional)'
    ]
    format_bullet_list(doc, backend)
    
    doc.add_heading('🎨  Frontend', 3)
    frontend = [
        'Vanilla HTML/CSS/JS (no framework)',
        'Chart.js — analytics charts',
        'Leaflet.js — interactive maps',
        'Font Awesome 6.4 — icons',
        'Google Fonts — Inter, Space Grotesk',
        'CSS custom properties (variables)',
        'localStorage for client state'
    ]
    format_bullet_list(doc, frontend)
    
    doc.add_heading('🚀  Deployment', 3)
    deployment = [
        'Docker (Dockerfile)',
        'Render.com (render.yaml)',
        'Supabase — managed PostgreSQL',
        'MongoDB Atlas — managed NoSQL',
        'AWS S3 — file storage',
        'Environment-based configuration',
        'Graceful fallback to dev mode'
    ]
    format_bullet_list(doc, deployment)
    
    doc.add_heading('System Architecture', 3)
    arch = doc.add_paragraph(
        'Client Browser → Express.js API → PostgreSQL / MongoDB\n'
        'Additional: express-rate-limit • Session Map with auto-cleanup • '
        'SQL migrations • Custom logger • Input validators • Error handling middleware'
    )
    arch.paragraph_format.space_before = Pt(6)
    arch.paragraph_format.space_after = Pt(6)
    
    add_section_break(doc)
    
    # ===== SLIDE 14: DATA EXPORT & REPORTING =====
    doc.add_heading('Data Export & Reporting', 1)
    doc.add_heading('CSV Exports, Audit Logs & Metrics', 2)
    
    doc.add_heading('Issues Export', 3)
    doc.add_paragraph('Access: PA + MCA')
    doc.add_paragraph('Endpoint: /api/admin/export/issues')
    doc.add_paragraph('Fields: Ticket ID, category, status, message, phone, source, dates')
    
    doc.add_heading('Bursary Export', 3)
    doc.add_paragraph('Access: MCA Only')
    doc.add_paragraph('Endpoint: /api/admin/export/bursaries')
    doc.add_paragraph('Fields: Ref code, student, institution, amount, status, decision')
    
    doc.add_heading('Constituents Export', 3)
    doc.add_paragraph('Access: MCA Only')
    doc.add_paragraph('Endpoint: /api/admin/export/constituents')
    doc.add_paragraph('Fields: Phone, national ID, name, location, village, verification')
    
    doc.add_heading('USSD Interactions', 3)
    doc.add_paragraph('Access: PA + MCA')
    doc.add_paragraph('Endpoint: /api/admin/export/ussd')
    doc.add_paragraph('Fields: Phone, input text, response, ref code, IP, timestamp')
    
    doc.add_heading('Audit Logs', 3)
    doc.add_paragraph('Access: PA + MCA')
    doc.add_paragraph('Via audit-trail module')
    doc.add_paragraph('Fields: Action type, user, role, IP, timestamp, description')
    
    doc.add_paragraph()
    doc.add_paragraph('All exports use UTF-8 BOM encoding for Excel compatibility • Date range filtering available on all exports')
    
    add_section_break(doc)
    
    # ===== SLIDE 15: SUMMARY & ROADMAP =====
    doc.add_heading('12. Summary & Roadmap', 1)
    doc.add_heading('Key Takeaways & Future Development Plans', 2)
    
    doc.add_heading('✅  Key Takeaways', 3)
    takeaways = [
        'Complete ward management dashboard — 6+ modules',
        'Multi-channel access: Web, USSD, Mobile App, WhatsApp',
        'Role-based access control (MCA, PA, Clerk)',
        'Real-time analytics with Chart.js & Leaflet.js maps',
        'Comprehensive audit trail & export capabilities',
        'Modern glassmorphism UI with responsive design',
        'AI assistant (Mai) for in-dashboard help',
        'Enterprise-grade security: bcrypt, rate limiting, CSP',
        'Graceful degradation — works even when DB is offline'
    ]
    format_bullet_list(doc, takeaways)
    
    doc.add_heading('🚀  Future Roadmap', 3)
    roadmap = [
        'Push notifications (Firebase Cloud Messaging)',
        'SMS bulk notifications (Africa\'s Talking API)',
        'Geographic issue heatmap visualization',
        'AI-powered issue categorization',
        'Multilingual support (Swahili/English)',
        'WebSocket real-time updates (replace polling)',
        'Unit, integration & E2E testing',
        'App store submission (iOS/Android)',
        'USSD shortcode registration'
    ]
    format_bullet_list(doc, roadmap)
    
    add_section_break(doc)
    
    # ===== SLIDE 16: CLOSING =====
    closing = doc.add_heading('Thank You', 0)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle2 = doc.add_paragraph('VOO Ward Admin Dashboard')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2_format = subtitle2.runs[0]
    subtitle2_format.bold = True
    subtitle2_format.size = Pt(14)
    
    tagline2 = doc.add_paragraph('Empowering Kyamatu Ward Through Digital Governance')
    tagline2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('Questions & Discussion')
    
    footer = doc.add_paragraph('VOO Ward Admin Dashboard — Empowering Local Governance Through Technology')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_format = footer.runs[0]
    footer_format.size = Pt(9)
    footer_format.italic = True
    
    # Save document
    output_file = r"c:\Users\kivuv\Documents\voo-ward-ussd\docs\VOO_Ward_Admin_Dashboard_Detailed_Proposal.docx"
    doc.save(output_file)
    return output_file

if __name__ == "__main__":
    output = create_detailed_docx()
    print(f"✓ Created detailed professional document: {output}")
